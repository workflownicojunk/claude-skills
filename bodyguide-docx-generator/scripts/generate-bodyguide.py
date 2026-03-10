#!/usr/bin/env python3
"""
BodyGuide DOCX Generator v4
Repliziert das echte Google Docs BodyGuide-Design 1:1.
Generiert Rezeptbilder via Nano Banana Pro Preview (Google Gemini).
13 statische Brand-Bilder (Nicol-Fotos, Illustrationen) eingebettet.

Usage:
    python3 generate-bodyguide.py input.json [output.docx]
    python3 generate-bodyguide.py --test
    python3 generate-bodyguide.py --test --no-images
"""

import json
import sys
import os
import base64
import time
import tempfile
from pathlib import Path
from io import BytesIO
import urllib.request

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === STATIC ASSETS ===
STATIC_IMAGES_DIR = Path(os.path.expanduser(
    '~/Desktop/Area/Produkte/body-guide/data/templates/static-images'
))

# === BRAND CONSTANTS (from real BodyGuide PDF) ===

COLORS = {
    'dark': RGBColor(0x24, 0x14, 0x04),      # Body text, dark brown
    'blue': RGBColor(0x7A, 0x94, 0xCC),      # Headlines, accent
    'brown': RGBColor(0x96, 0x82, 0x7D),     # Captions, secondary
    'cream': RGBColor(0xFA, 0xED, 0xE1),     # Page background
    'maroon': RGBColor(0x8B, 0x1A, 0x1A),    # Table headers (Tipps)
    'white': RGBColor(0xFF, 0xFF, 0xFF),
}

FONT_HEADLINE = 'Playfair Display'
FONT_BODY = 'Noto Sans'

# Gemini image prompt prefix for food photography
FOOD_PHOTO_PREFIX = (
    "Editorial food photography, warm cream tones, soft natural lighting from left, "
    "shallow depth of field, overhead angle 30-45 degrees, rustic ceramic plate or bowl, "
    "premium magazine aesthetic, NOT stock photo, NOT clinical, NOT generic AI aesthetic. "
    "Photo of: "
)

# === HELPERS ===

def set_font(run, name=FONT_BODY, size=Pt(11), color=None, bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), name)


def set_page_bg(doc, color_hex='FAEDE1'):
    """Set page background color for the entire document."""
    bg = parse_xml(
        f'<w:background {nsdecls("w")} w:color="{color_hex}" w:themeColor="background1"/>'
    )
    doc.element.insert(0, bg)
    # Also set displayBackgroundShape
    settings = doc.settings.element
    display_bg = parse_xml(f'<w:displayBackgroundShape {nsdecls("w")}/>')
    settings.append(display_bg)


def add_headline(doc, text, level=1, color=None):
    sizes = {1: Pt(28), 2: Pt(22), 3: Pt(18)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, FONT_HEADLINE, sizes.get(level, Pt(18)),
             color or COLORS['blue'], bold=True, italic=True)
    p.space_after = Pt(8)
    p.space_before = Pt(16) if level > 1 else Pt(0)
    return p


def add_body(doc, text, color=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, FONT_BODY, Pt(11), color or COLORS['dark'], bold=bold)
    p.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bold_body(doc, text, color=None):
    return add_body(doc, text, color or COLORS['dark'], bold=True)


def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_font(run, FONT_BODY, Pt(11), color or COLORS['dark'])
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, font_name=FONT_BODY, size=Pt(11), color=None,
                  bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    set_font(run, font_name, size, color or COLORS['dark'], bold=bold, italic=italic)


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with value (sz, color, val)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, props in kwargs.items():
        elem = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{props.get("val", "single")}" '
            f'w:sz="{props.get("sz", "4")}" w:space="0" '
            f'w:color="{props.get("color", "auto")}"/>'
        )
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def hide_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def generate_food_image(recipe_name, ingredients, api_key):
    """Generate food image via Gemini 2.5 Flash Image."""
    prompt = (
        f"{FOOD_PHOTO_PREFIX}{recipe_name}. "
        f"Main ingredients visible: {ingredients[:200]}. "
        "Beautiful plating on rustic surface."
    )
    url = f"https://generativelanguage.googleapis.com/v1beta/models/nano-banana-pro-preview:generateContent?key={api_key}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"responseModalities": ["IMAGE", "TEXT"]}
    }
    data = json.dumps(payload).encode()
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    try:
        resp = urllib.request.urlopen(req, timeout=90)
        result = json.loads(resp.read())
        for c in result.get("candidates", []):
            for p in c.get("content", {}).get("parts", []):
                if "inlineData" in p:
                    return BytesIO(base64.b64decode(p["inlineData"]["data"]))
    except Exception as e:
        print(f"    [WARN] Bildgenerierung fehlgeschlagen: {recipe_name} ({e})")
    return None


def add_static_image(doc, filename, width=Cm(14), center=True):
    """Insert a static brand image from the assets directory."""
    path = STATIC_IMAGES_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=width)
        if center:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f"    [WARN] Statisches Bild nicht gefunden: {path}")
        return False


def download_image(url, timeout=15):
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'BodyGuide/2.0'})
        return BytesIO(urllib.request.urlopen(req, timeout=timeout).read())
    except Exception as e:
        print(f"    [WARN] Download fehlgeschlagen: {url} ({e})")
        return None


# === DOCUMENT SECTIONS ===

def build_title_page(doc, data):
    customer = data['customer']
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{customer['firstName']}'s")
    set_font(run, FONT_HEADLINE, Pt(32), COLORS['blue'], bold=True, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Personal Body Guide')
    set_font(run, FONT_HEADLINE, Pt(40), COLORS['dark'], bold=True, italic=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run(
        'Dein Komplettsystem f\u00fcr Ern\u00e4hrung und Training -\n'
        'f\u00fcr einen fitten, festen und definierten K\u00f6rper!'
    )
    set_font(run, FONT_HEADLINE, Pt(14), COLORS['dark'], bold=True, italic=True)

    # Nicol Cover-Foto
    add_static_image(doc, 'cover-nicol.jpg', width=Cm(12))

    doc.add_page_break()


def build_intro(doc, data):
    customer = data['customer']

    p = doc.add_paragraph()
    run = p.add_run(f"Liebe {customer['firstName']}")
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    add_body(doc, '')
    add_body(doc, 'geh\u00f6rst du auch zu den Frauen, die schon alles versucht haben - aber einfach '
             'nicht weiter abnehmen?')
    add_body(doc, 'Vielleicht hast du aus lauter Verzweiflung schon')

    for b in ['gefastet,', 'gehungert,', 'zig Di\u00e4ten gemacht,',
              'trainiert wie eine Verr\u00fcckte',
              'und wochenlang auf S\u00fc\u00dfes verzichtet?']:
        add_bullet(doc, b)

    add_body(doc, 'Ich sag dir was:')
    add_bold_body(doc, 'Diese Qualen sind jetzt vorbei (wenn du willst, f\u00fcr immer)')

    add_body(doc, '')
    add_body(doc, 'Du erinnerst dich bestimmt an die vielen Fragen, die du f\u00fcr diesen Guide beantwortet hast.')
    add_body(doc, 'Basierend auf deinen Angaben bekommst du jetzt einen Ern\u00e4hrungs- und '
             'Trainingsplan, der seinesgleichen sucht.')

    add_body(doc, '')
    add_headline(doc, 'Warum ist dieser Plan so einzigartig?', level=2)
    add_body(doc, 'Weil es die Rezepte, die du auf den n\u00e4chsten Seiten findest, in dieser '
             'Zusammenstellung und mit diesen Mengenangaben nur ein einziges Mal gibt - '
             'und zwar f\u00fcr dich.')
    add_body(doc, 'Dein pers\u00f6nlicher BodyGuide ist nichts von der Stange, sondern '
             'ausschlie\u00dflich auf deinen K\u00f6rper zugeschnitten.')
    add_body(doc, 'Es steckt quasi deine DNA drin.')

    doc.add_page_break()

    # Baukastensystem
    add_headline(doc, 'Wie ist der Plan aufgebaut?', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Dein BodyGuide arbeitet mit einem ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('flexiblen Baukastensystem:')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)

    for b in ['5 Bausteine (A bis E) mit je drei Mahlzeiten\n(Genaueres dazu erf\u00e4hrst du gleich)',
              'frei kombinierbar nach deinem (Essens-)Rhythmus',
              'anpassbar an deinen Alltag zwischen Job, Familie und Freizeit',
              'wachsend mit deinen Fortschritten']:
        add_bullet(doc, b)

    add_body(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Ganz wichtig:')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run.underline = True
    run = p.add_run(' Du musst nicht alles sofort umsetzen!')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    add_bold_body(doc, 'Starte mit einem Baustein, der dir am meisten zusagt.')
    add_body(doc, 'Nach etwa 4 Wochen kennst du alle Optionen und wei\u00dft genau, was dir am besten liegt.')
    add_body(doc, 'Danach f\u00e4llt es dir immer leichter, dranzubleiben.')

    # Baukasten-Grafik
    add_static_image(doc, 'baukasten-grafik.jpg', width=Cm(10))

    doc.add_page_break()

    # Sweet Spot
    add_headline(doc, 'Was ist das Besondere am BodyGuide?', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Er basiert auf dem ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('"Sweet Spot"-Prinzip')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run = p.add_run(' - der Punkt, an dem dein Stoffwechsel optimal arbeitet.')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    add_body(doc, '')
    add_bold_body(doc, 'Und dann passiert Folgendes:')
    add_body(doc, 'Die l\u00e4stigen Fettpolster, die dich schon seit Jahren st\u00f6ren, verschwinden.')
    add_body(doc, 'Nicht irgendwo.')
    add_body(doc, 'Sondern genau an deinen Problemzonen.')

    # Nicol Spiegel-Foto
    add_static_image(doc, 'nicol-spiegel.jpg', width=Cm(10))

    add_body(doc, '')
    add_headline(doc, 'Warum funktioniert das so gut?', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Weil dein Stoffwechsel einen nat\u00fcrlichen Rhythmus hat, \u00e4hnlich wie eine innere Uhr.')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)

    add_body(doc, 'Der BodyGuide ist auf die innere Uhr deines K\u00f6rper abgestimmt und basiert auf '
             '3 wissenschaftlich fundierten S\u00e4ulen:')

    doc.add_page_break()

    # 3 Saeulen
    pillar_images = {
        '1': 'sweet-spot-illustration.jpg',
        '2': 'nicol-training.jpg',
        '3': 'training-kombination-kreise.jpg',
    }
    for num, title, texts in [
        ('1', 'Optimales Timing', [
            ('Dein ', False), ('Stoffwechsel arbeitet unterschiedlich effektiv', True),
            (' - je nach Tageszeit.', False),
        ]),
        ('2', 'Aktivit\u00e4tsbasierte Ern\u00e4hrung', [
            ('Bevor du dein Workout machst, braucht dein K\u00f6rper Energie. Die bekommst du \u00fcber ', False),
            ('Kohlenhydrate.', True),
        ]),
        ('3', 'Effektive Trainings-Kombination', [
            ('', False), ('Krafttraining', True), (' erh\u00f6ht deinen Grundumsatz.', False),
        ]),
    ]:
        add_headline(doc, f'#{num}) {title}', level=2)
        if num in pillar_images:
            add_static_image(doc, pillar_images[num], width=Cm(10))
        p = doc.add_paragraph()
        for text, is_bold in texts:
            if text:
                run = p.add_run(text)
                set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=is_bold)
        add_body(doc, '')
        doc.add_page_break()


def build_fitness_journey(doc, data):
    customer = data['customer']
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(36)
    run = p.add_run('Deine Fitness Journey')
    set_font(run, FONT_HEADLINE, Pt(24), COLORS['blue'], bold=True, italic=True)

    fields = [
        ('Name', customer.get('name', customer['firstName'])),
        ('Alter', customer.get('alter', '')),
        ('Gr\u00f6\u00dfe (cm)', customer.get('groesse', '')),
        ('Aktuelles Gewicht (kg)', customer.get('aktuellesGewicht', '')),
        ('Zielgewicht (kg)', customer.get('zielgewicht', '')),
        ('Gesamtenergieumsatz', customer.get('gesamtenergieumsatz', '')),
        ('Tagesanpassung pro Tag', customer.get('tagesanpassung', '')),
        ('Angepasster Energieumsatz', customer.get('angepassterEnergieumsatz', '')),
        ('Ern\u00e4hrungsform', customer.get('ernaehrungsform', '')),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hide_table_borders(table)

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.height = Cm(1.5)

        set_cell_text(row.cells[0], label, FONT_HEADLINE, Pt(12),
                      COLORS['blue'], bold=True, italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], str(value), FONT_HEADLINE, Pt(12),
                      COLORS['blue'], bold=True, italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add vertical divider between columns
        set_cell_border(row.cells[0], right={"sz": "4", "color": "C0C0C0", "val": "single"})

    doc.add_page_break()


def build_meal_planner(doc, data):
    meal_plan = data.get('mealPlan', {})

    add_headline(doc, 'Dein Weekly Meal Planner', level=1)

    p = doc.add_paragraph()
    run = p.add_run('In dem nun folgenden Weekly Meal Planner findest du die bereits erw\u00e4hnten ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('5 Bausteine (A - E).')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)

    add_body(doc, 'Ein einzelner Baustein deckt jeweils einen Tag ab - mit Fr\u00fchst\u00fcck, '
             'Mittag- und Abendessen.')

    p = doc.add_paragraph()
    run = p.add_run('Du kannst mit diesem Plan ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('sehr flexibel')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run = p.add_run(' umgehen.')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    add_body(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Was ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run = p.add_run('nicht')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run.underline = True
    run = p.add_run(' funktioniert:')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)

    add_body(doc, '3x hintereinander das Fr\u00fchst\u00fccks-, Mittag- oder Abendrezept zubereiten.')

    add_bold_body(doc, 'Du brauchst also immer diese Reihenfolge:', COLORS['dark'])
    for item in ['Energiestart', 'Energie-Power', 'Ausklang']:
        add_bullet(doc, item)

    doc.add_page_break()

    # Meal Planner Table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('WEEKLY MEAL\nPLANNER')
    set_font(run, FONT_HEADLINE, Pt(28), COLORS['blue'], bold=True, italic=False)
    run = p.add_run(f'  Name\n  {data["customer"].get("name", data["customer"]["firstName"])}')
    set_font(run, FONT_HEADLINE, Pt(14), COLORS['blue'], bold=True, italic=True)

    add_body(doc, '')

    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Option', 'Energiestart', 'Energie-Power', 'Ausklang']
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, FONT_HEADLINE, Pt(11),
                      COLORS['dark'], bold=True, italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, opt in enumerate(['A', 'B', 'C', 'D', 'E'], start=1):
        mp = meal_plan.get(opt, {})
        set_cell_text(table.rows[i].cells[0], opt, FONT_HEADLINE, Pt(14),
                      COLORS['blue'], bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for j, key in enumerate(['meal1', 'meal2', 'meal3'], start=1):
            set_cell_text(table.rows[i].cells[j], mp.get(key, ''),
                          FONT_BODY, Pt(10), COLORS['dark'],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Dashed red borders between cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "2", "color": "CC0000", "val": "dashed"},
                bottom={"sz": "2", "color": "CC0000", "val": "dashed"},
                left={"sz": "2", "color": "CC0000", "val": "dashed"},
                right={"sz": "2", "color": "CC0000", "val": "dashed"})

    doc.add_page_break()


def build_recipe_page(doc, data, option, meal_num, recipe, api_key=None, gen_images=True):
    """Build a single recipe page with 2-column layout matching the real BodyGuide."""

    # 2-column layout via table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hide_table_borders(table)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # Set column widths (40% left, 60% right)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # === LEFT COLUMN ===
    # Baustein label
    p = left_cell.paragraphs[0]
    run = p.add_run(f'Baustein {option}')
    set_font(run, FONT_BODY, Pt(11), COLORS['brown'])

    # Recipe name
    p = left_cell.add_paragraph()
    run = p.add_run(recipe.get('name', f'Rezept {option}.{meal_num}'))
    set_font(run, FONT_HEADLINE, Pt(24), COLORS['blue'], bold=True, italic=True)

    # Recipe image
    img_data = None
    images = data.get('images', {}).get('recipeImages', {})
    img_key = f'{option}.{meal_num}'

    if img_key in images and images[img_key]:
        img_url = images[img_key]
        if img_url.startswith('http'):
            img_data = download_image(img_url)
        elif os.path.exists(img_url):
            with open(img_url, 'rb') as f:
                img_data = BytesIO(f.read())
    elif gen_images and api_key:
        ingredients_str = recipe.get('zutaten', '')
        if isinstance(ingredients_str, list):
            ingredients_str = ', '.join(ingredients_str)
        img_data = generate_food_image(recipe.get('name', ''), ingredients_str, api_key)
        if img_data:
            # Cache the generated image
            cache_dir = Path(tempfile.gettempdir()) / 'bodyguide-images'
            cache_dir.mkdir(exist_ok=True)
            cache_path = cache_dir / f'{option}_{meal_num}.png'
            with open(cache_path, 'wb') as f:
                f.write(img_data.getvalue())
            img_data.seek(0)
            time.sleep(2)  # Rate limit

    if img_data:
        p = left_cell.add_paragraph()
        run = p.add_run()
        run.add_picture(img_data, width=Cm(6.5))

    # Kalorien
    p = left_cell.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run('Kalorien')
    set_font(run, FONT_BODY, Pt(10), COLORS['brown'])
    p = left_cell.add_paragraph()
    run = p.add_run(f"{recipe.get('kalorien', '---')} Kalorien")
    set_font(run, FONT_BODY, Pt(12), COLORS['dark'], bold=True)

    # === RIGHT COLUMN ===
    # Zutaten
    p = right_cell.paragraphs[0]
    run = p.add_run('Zutaten')
    set_font(run, FONT_HEADLINE, Pt(18), COLORS['blue'], bold=True, italic=True)

    zutaten = recipe.get('zutaten', '')
    items = zutaten if isinstance(zutaten, list) else [z.strip().lstrip('*-').strip()
            for z in str(zutaten).split('\n') if z.strip()]
    for item in items:
        p = right_cell.add_paragraph()
        run = p.add_run(f'\u2022  {item}')
        set_font(run, FONT_BODY, Pt(10), COLORS['dark'])
        p.space_after = Pt(1)

    # Zubereitung
    p = right_cell.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run('Zubereitung')
    set_font(run, FONT_HEADLINE, Pt(18), COLORS['blue'], bold=True, italic=True)

    zubereitung = recipe.get('zubereitung', '')
    steps = zubereitung if isinstance(zubereitung, list) else [s.strip().lstrip('*-').strip()
            for s in str(zubereitung).split('\n') if s.strip()]
    for step in steps:
        p = right_cell.add_paragraph()
        run = p.add_run(f'\u2022  {step}')
        set_font(run, FONT_BODY, Pt(10), COLORS['dark'])
        p.space_after = Pt(1)

    # CravingControl
    cc = recipe.get('cravingControl', {})
    if cc:
        p = right_cell.add_paragraph()
        p.space_before = Pt(16)
        run = p.add_run('Craving Control Strategie')
        set_font(run, FONT_HEADLINE, Pt(16), COLORS['blue'], bold=True, italic=True)

        for label, key in [('Was', 'was'), ('Warum', 'warum'), ('L\u00f6sung', 'loesung')]:
            p = right_cell.add_paragraph()
            run = p.add_run(label)
            set_font(run, FONT_HEADLINE, Pt(11), COLORS['blue'], bold=True, italic=True)
            p = right_cell.add_paragraph()
            run = p.add_run(cc.get(key, ''))
            set_font(run, FONT_BODY, Pt(10), COLORS['dark'])
            p.space_after = Pt(2)

    doc.add_page_break()


def build_all_recipes(doc, data, api_key=None, gen_images=True):
    recipes = data.get('recipes', {})
    total = sum(len(recipes.get(opt, [])) for opt in 'ABCDE')
    current = 0
    for option in ['A', 'B', 'C', 'D', 'E']:
        option_recipes = recipes.get(option, [])
        for meal_num in range(1, 4):
            idx = meal_num - 1
            current += 1
            if idx < len(option_recipes):
                recipe = option_recipes[idx]
            else:
                recipe = {'name': f'Rezept {option}.{meal_num}', 'kalorien': '---',
                          'zutaten': '', 'zubereitung': '', 'cravingControl': {}}
            print(f"    Rezept {current}/15: {recipe.get('name', '?')}")
            build_recipe_page(doc, data, option, meal_num, recipe, api_key, gen_images)


def build_cheat_days(doc):
    add_headline(doc, 'Was ist mit Tag 6 und 7 in der Woche?', level=1)
    add_body(doc, 'Jetzt wunderst du dich vielleicht:')
    add_body(doc, '"Das sind ja nur 5 Tage pro Woche!"')
    add_body(doc, 'Gut aufgepasst.')
    add_bold_body(doc, 'Und jetzt habe ich eine prima Nachricht f\u00fcr dich:')
    p = doc.add_paragraph()
    run = p.add_run('Die anderen 2 Tage sind ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('Cheat-Days')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run.underline = True
    run = p.add_run(' (cheat = schummeln).')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    add_body(doc, 'An diesen 2 Tagen (beispielsweise am Wochenende) musst du dich nicht an den Plan halten.')
    add_body(doc, 'Du darfst, wenn du m\u00f6chtest, mehr Kohlenhydrate essen.')
    add_body(doc, 'Die Wirkung?')

    for b in ['Dein K\u00f6rper ist zufrieden.',
              'Du hast weniger Stress und kannst loslassen.',
              'Dein K\u00f6rper bekommt nicht das Gef\u00fchl, dass du ihm etwas wegnimmst.']:
        add_bullet(doc, b)

    add_body(doc, '')
    add_bold_body(doc, 'Er r\u00fcckt kein Gramm Fett mehr raus.', COLORS['dark'])
    add_body(doc, 'Und stellt auf stur.')

    doc.add_page_break()

    add_headline(doc, 'Was du an den Cheat-Days nicht machen darfst', level=2)
    add_body(doc, 'Dich vollstopfen.')
    add_body(doc, 'Komplett \u00fcber die Str\u00e4nge schlagen.')
    add_body(doc, 'Binge Eating.')
    add_body(doc, 'Das geht nat\u00fcrlich nicht.')
    add_body(doc, 'Aber wenn du es \u00fcbertreibst, k\u00f6nnen die n\u00e4chsten 5 Tage das nat\u00fcrlich nicht mehr wettmachen.')

    doc.add_page_break()


def build_drinks_guide(doc):
    add_headline(doc, 'Die richtigen Getr\u00e4nke w\u00e4hlen', level=1)
    add_body(doc, 'Dieses Kapitel des BodyGuides ist extrem wichtig f\u00fcr dich.')
    add_body(doc, 'Das Thema Getr\u00e4nke f\u00e4llt gerne mal unter den Tisch.')
    add_bold_body(doc, 'Aber wenn du erfolgreich Fett verbrennen und einen definierten Body '
                  'bekommen m\u00f6chtest, darfst du deine Hydration nicht dem Zufall \u00fcberlassen.')

    # Getränke Shaker-Foto
    add_static_image(doc, 'getraenke-shaker.jpg', width=Cm(12))

    add_body(doc, '')
    add_headline(doc, '3 wissenschaftlich belegte Fakten', level=2)
    add_bold_body(doc, '#1) Bei einer Dehydration von nur 2 %, reduziert sich deine '
                  'Stoffwechselleistung um bis zu 30 %.')
    add_body(doc, 'Das belegt eine Studie vom Journal of Clinical Endocrinology & Metabolism von 2003.')

    add_bold_body(doc, '#2) Wasser erh\u00f6ht den Energieverbrauch um 24-30 % f\u00fcr die n\u00e4chsten '
                  '60 Minuten nach dem Trinken (Thermogenese-Effekt).')

    add_bold_body(doc, '#3) Jeder Liter kaltes Wasser verbrennt zus\u00e4tzlich 35-40 kcal durch die '
                  'Erw\u00e4rmung auf K\u00f6rpertemperatur.')

    add_body(doc, '')
    add_headline(doc, 'So berechnest du deine optimale Trinkmenge', level=2)
    add_bold_body(doc, '35 ml pro kg K\u00f6rpergewicht + 500-750 ml pro Trainingsstunde.')

    for b in ['2,3 Liter Basis', '500-750 ml zus\u00e4tzlich bei einer Stunde Workout',
              'zus\u00e4tzlich 500 ml pro 5\u00b0C \u00fcber 25\u00b0C']:
        add_bullet(doc, b)

    doc.add_page_break()

    add_headline(doc, 'Was du trinken solltest', level=2)
    add_body(doc, 'Allem voran (du ahnst es bereits): Wasser.')
    add_bold_body(doc, 'Wasser ist kein "Nice-to-have"!')
    for b in ['Gr\u00fcner Tee', 'Ingwer-Zitronen-Wasser']:
        add_bullet(doc, b, COLORS['dark'])
    add_body(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Schwarzer Kaffee')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run = p.add_run(' kurbelt ebenfalls die Fettverbrennung an.')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    add_body(doc, 'Trinke hiervon aber maximal 4 Tassen pro Tag.')

    # Zitronenwasser-Foto
    add_static_image(doc, 'getraenke-zitronenwasser.jpg', width=Cm(10))

    doc.add_page_break()

    add_headline(doc, 'Was du nicht trinken solltest', level=2)
    add_bold_body(doc, '1. Fruchts\u00e4fte')
    add_body(doc, 'Nicht einmal 100 % Direktsaft.')
    add_body(doc, '')
    add_bold_body(doc, '2. Smoothies (Fertigprodukte)')
    add_body(doc, 'Auch hier steckt viel Fruktose drin.')
    add_body(doc, '')
    add_bold_body(doc, '3. "Zero"-Getr\u00e4nke')
    add_body(doc, 'Die enthaltenen S\u00fc\u00dfstoffe k\u00f6nnen dein Insulin beeinflussen.')

    add_body(doc, '')
    add_headline(doc, 'Alkohol: Ja oder nein?', level=2)
    add_bold_body(doc, 'Immer, wenn du Alkohol trinkst, ist das wie eine Pause-Taste f\u00fcr deine '
                  'Body-Transformation, die dich 2-3 Tage Fortschritt kostet.')
    add_bold_body(doc, 'Trinke nur 1x pro Woche Alkohol.')
    for b in ['Ein Glas Wei\u00dfweinschorle vor dem Essen.',
              'Oder ein Glas Rotwein w\u00e4hrend einer proteinreichen Mahlzeit.']:
        add_bullet(doc, b)
    add_bold_body(doc, 'Bier ist leider keine Option.')

    add_body(doc, '')
    add_headline(doc, 'Die praktische Umsetzung f\u00fcr dich', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Mach es dir zur Gewohnheit, ein ')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])
    run = p.add_run('t\u00e4gliches Trink-Protokoll')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'], bold=True)
    run = p.add_run(' einzuf\u00fchren.')
    set_font(run, FONT_BODY, Pt(11), COLORS['dark'])

    schedule = [
        '6:00 Uhr \u00bb 500 ml lauwarmes Wasser',
        '7:30 Uhr \u00bb 250 ml Wasser vor dem Fr\u00fchst\u00fcck',
        '10:00 Uhr \u00bb 300 ml Gr\u00fcner Tee',
        '12:00 Uhr \u00bb 250 ml Wasser vor dem Mittagessen',
        '15:00 Uhr \u00bb 300 ml Wasser',
        '18:30 Uhr \u00bb 250 ml vor dem Abendessen',
        '21:00 Uhr \u00bb 200 ml Kr\u00e4utertee',
    ]
    for s in schedule:
        add_bullet(doc, s)

    # Trink-Protokoll Illustration
    add_static_image(doc, 'trink-protokoll-illustration.jpg', width=Cm(10))

    doc.add_page_break()


def build_snack_section(doc, data):
    add_headline(doc, 'Clever snacken', level=1)
    add_body(doc, 'Ja, du darfst snacken.')
    add_body(doc, 'Ein Snack pro Tag (100\u2013150 kcal) passt perfekt in dein Defizit.')

    add_headline(doc, 'Wann darf ich die Snacks essen?', level=2)
    add_body(doc, 'Entweder am Nachmittag. Oder kurz vor dem Schlafengehen.')
    add_body(doc, 'Bitte esse Snacks zu keinem anderen Zeitpunkt!')

    snacks = data.get('snacks', {})
    categories = [
        ('VEGAN', 'vegan'), ('VEGGIE', 'vegetarian'),
        ('OMNIVOR', 'flexitarian'), ('NACHT-REGENERATION', 'midnight'),
    ]

    for cat_label, cat_key in categories:
        items = snacks.get(cat_key, [])
        if not items:
            continue

        add_body(doc, '')
        table = doc.add_table(rows=len(items) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['kcal', cat_label, 'Zutaten', 'Zubereitung']
        for j, h in enumerate(headers):
            set_cell_text(table.rows[0].cells[j], h, FONT_HEADLINE, Pt(9),
                          COLORS['white'], bold=True, italic=True)
            set_cell_shading(table.rows[0].cells[j], '8B1A1A')

        for i, snack in enumerate(items, 1):
            set_cell_text(table.rows[i].cells[0], '~130', FONT_BODY, Pt(9), COLORS['dark'], bold=True)
            set_cell_text(table.rows[i].cells[1], snack.get('name', ''), FONT_BODY, Pt(9), COLORS['dark'])
            set_cell_text(table.rows[i].cells[2], snack.get('ingredients', ''), FONT_BODY, Pt(8), COLORS['dark'])
            set_cell_text(table.rows[i].cells[3], snack.get('preparation', ''), FONT_BODY, Pt(8), COLORS['dark'])

    doc.add_page_break()


def build_workout_and_closing(doc, data):
    add_headline(doc, 'Workout-Guide', level=1)

    # Nicol Workout-Foto
    add_static_image(doc, 'nicol-workout.jpg', width=Cm(8))

    add_body(doc, 'Dein Workout-Guide besteht aus 3 effektiven Home-Workouts + 1 Bonus-Workout.')

    add_headline(doc, 'So machst du die Workouts', level=2)
    for s in ['Je \u00dcbung: 30 Sek pro Seite', 'Pause: 1 Minute zwischen den \u00dcbungen',
              'Ziel: 3 Durchg\u00e4nge', 'Trainiere 3-4x die Woche']:
        add_bullet(doc, s)

    add_headline(doc, 'Deine Workout-Videos', level=2)
    for i, w in enumerate(['Triceps Training - Dips', 'Hundreds - Bauchworkout',
                           'Beine Jump Squats', 'Mountainclimber Optionen'], 1):
        add_body(doc, f'{i}.: {w}')

    doc.add_page_break()

    # Partner Links
    add_headline(doc, 'Deine Link-Sammlung f\u00fcr die besten Supplements und Tools', level=1)
    partners = [
        ('HARVEST REPUBLIC', 'NICOL', 'Premium Protein-Shakes'),
        ('LES MILLS EQUIPMENT', 'NICOL10', 'Hochwertiges Home-Equipment'),
        ('INNONATURE', 'NICOL10', 'Omega-3, Lemon Morning und mehr'),
    ]
    for name, code, desc in partners:
        add_headline(doc, name, level=3)
        add_body(doc, desc)
        add_bold_body(doc, f'Der Code lautet: {code}')
        add_body(doc, '')

    doc.add_page_break()

    # Abschluss
    add_headline(doc, 'Ein paar Worte zum Abschluss', level=1)
    add_body(doc, 'Meine Wahrheit ist:')
    add_bold_body(doc, 'Mein K\u00f6rper, meine Entscheidung.')
    add_body(doc, 'Do it or lose it.')
    add_body(doc, 'Was ist deine Wahrheit?')
    add_body(doc, 'Ready, dir den K\u00f6rper zu holen, den du schon immer wolltest?')

    p = doc.add_paragraph()
    run = p.add_run("Then let's do this!")
    set_font(run, FONT_HEADLINE, Pt(18), COLORS['blue'], bold=True, italic=True)

    add_body(doc, '')
    p = doc.add_paragraph()
    run = p.add_run('Deine Nicol')
    set_font(run, FONT_HEADLINE, Pt(16), COLORS['dark'], bold=True, italic=True)


# === MAIN ===

def generate_bodyguide(data, output_path, api_key=None, gen_images=True):
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cream background
    set_page_bg(doc, 'FAEDE1')

    # Default font
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS['dark']

    print("  [1/9] Titelseite...")
    build_title_page(doc, data)
    print("  [2/9] Einleitung...")
    build_intro(doc, data)
    print("  [3/9] Fitness Journey...")
    build_fitness_journey(doc, data)
    print("  [4/9] Weekly Meal Planner...")
    build_meal_planner(doc, data)
    print("  [5/9] Rezepte (15 Mahlzeiten)...")
    build_all_recipes(doc, data, api_key, gen_images)
    print("  [6/9] Cheat Days...")
    build_cheat_days(doc)
    print("  [7/9] Getr\u00e4nke-Guide...")
    build_drinks_guide(doc)
    print("  [8/9] Snack-Karten...")
    build_snack_section(doc, data)
    print("  [9/9] Workout + Abschluss...")
    build_workout_and_closing(doc, data)

    doc.save(output_path)
    print(f"\nBodyGuide gespeichert: {output_path}")
    return output_path


def generate_test_data():
    recipes = {}
    sample_meals = {
        'A': [
            ('Beeriger Morgenzauber', '342', '150g laktosefreier MQ\n30g gf Haferflocken\n50g TK-Beeren\n10g Chia\n5g Lein\u00f6l\n5g Honig',
             'Alle Zutaten in Sch\u00fcssel geben\nChiasamen 5 Min quellen lassen\numr\u00fchren\ngenie\u00dfen'),
            ('Goldene Lachspfanne', '522', '130g Lachs roh\n150g Brokkoli\n100g Zucchini\n80g Paprika\n40g Avocado\n10g Oliven\u00f6l',
             'Gem\u00fcse in Pfanne/Airfryer 5 Min garen\nLachs w\u00fcrzen\n8 Min braten\nmit Avocado servieren'),
            ('Frische Thunfisch-Bowl', '378', '100g Thunfisch Dose\n100g lf H\u00fcttenk\u00e4se\n100g Gurke\n80g Tomaten\n50g Avocado\n5g Oliven\u00f6l',
             'Thunfisch abtropfen\nGem\u00fcse w\u00fcrfeln\nmit H\u00fcttenk\u00e4se und Avocado anrichten\n\u00d6l dar\u00fcber'),
        ],
        'B': [
            ('Sonniger Eierteller', '350', '2 Eier (120g roh)\n100g Tomaten\n80g Paprika\n50g Avocado\n5g Oliven\u00f6l\n1 Prise Salz',
             'Eier in \u00d6l als R\u00fchrei braten\nGem\u00fcse frisch dazu schneiden\nmit Avocado servieren'),
            ('Knusprige H\u00e4hnchenfreude', '536', '150g H\u00e4hnchenbrust roh\n60g Quinoa roh\n100g Brokkoli\n100g Zucchini\n80g Paprika\n10g Oliven\u00f6l',
             'Quinoa kochen (15 Min)\nH\u00e4hnchen w\u00fcrzen+in Pfanne/Airfryer 10 Min braten\nGem\u00fcse dazu garen'),
            ('Leichte Lachs-Scheiben', '375', '90g R\u00e4ucherlachs\n100g lf Skyr\n100g Gurke\n80g Tomaten\n50g Avocado\n3g Oliven\u00f6l',
             'Lachs auf Teller anrichten\nGurke+Tomaten w\u00fcrfeln\nSkyr als Dip\nAvocado dazu\n\u00d6l dr\u00fcber'),
        ],
    }
    for opt in ['C', 'D', 'E']:
        sample_meals[opt] = sample_meals['A']  # Reuse for test

    for opt, meals in sample_meals.items():
        recipes[opt] = []
        for name, kcal, zutaten, zubereitung in meals:
            recipes[opt].append({
                'name': name, 'kalorien': kcal,
                'zutaten': zutaten, 'zubereitung': zubereitung,
                'cravingControl': {
                    'was': '6:00 Morgenm\u00fcdigkeit nach 5h Schlaf',
                    'warum': 'Stress Level 8 ersch\u00f6pft Energiereserven fr\u00fch',
                    'loesung': 'Protein+Ballaststoffe halten Blutzucker stabil bis Mittag',
                },
            })

    snacks = {
        'vegan': [{'name': f'Veganer Snack {i}', 'ingredients': 'N\u00fcsse, Datteln, Kakao', 'preparation': 'Alles vermengen.'} for i in range(1, 7)],
        'vegetarian': [{'name': f'Veggie Snack {i}', 'ingredients': 'Quark, Beeren, Honig', 'preparation': 'Quark mit Beeren mischen.'} for i in range(1, 7)],
        'flexitarian': [{'name': f'Flexi Snack {i}', 'ingredients': 'Eier, Avocado', 'preparation': 'Ei kochen, Avocado aufschneiden.'} for i in range(1, 6)],
        'midnight': [{'name': f'Nacht-Snack {i}', 'ingredients': 'Casein-Quark, Zimt', 'preparation': 'Quark mit Zimt verr\u00fchren.'} for i in range(1, 6)],
    }

    return {
        'customer': {
            'firstName': 'Sara', 'name': 'Sara Thaler', 'alter': '53',
            'groesse': '160', 'aktuellesGewicht': '52', 'zielgewicht': '48',
            'gesamtenergieumsatz': '1586 kcal', 'tagesanpassung': '-336 kcal',
            'angepassterEnergieumsatz': '1250 kcal', 'ernaehrungsform': 'Low Carb',
        },
        'mealPlan': {
            opt: {
                'meal1': recipes[opt][0]['name'],
                'meal2': recipes[opt][1]['name'],
                'meal3': recipes[opt][2]['name'],
            } for opt in ['A', 'B', 'C', 'D', 'E']
        },
        'recipes': recipes,
        'snacks': snacks,
        'images': {'recipeImages': {}, 'heroImages': {}},
    }


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 generate-bodyguide.py input.json [output.docx]")
        print("       python3 generate-bodyguide.py --test [--no-images]")
        sys.exit(1)

    gen_images = '--no-images' not in sys.argv

    # Load API key
    api_key = None
    if gen_images:
        try:
            import subprocess
            result = subprocess.run(['grep', '^GOOGLE_API_KEY=', os.path.expanduser('~/Desktop/.env')],
                                    capture_output=True, text=True)
            if result.stdout:
                api_key = result.stdout.strip().split('=', 1)[1]
                print(f"  Gemini API Key geladen (Bildgenerierung aktiv)")
        except Exception:
            pass
        if not api_key:
            print("  [WARN] GOOGLE_API_KEY nicht gefunden, Bildgenerierung deaktiviert")
            gen_images = False

    if '--test' in sys.argv:
        data = generate_test_data()
        output = '/tmp/bodyguide-test-v2.docx'
    else:
        input_file = [a for a in sys.argv[1:] if not a.startswith('--')][0]
        with open(input_file) as f:
            data = json.load(f)
        output = sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith('--') else input_file.replace('.json', '.docx')

    generate_bodyguide(data, output, api_key, gen_images)
